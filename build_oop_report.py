from __future__ import annotations

from datetime import datetime
from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"D:\PROJECT\PACMAN OOP")
REPORTS = ROOT / "reports"


def latex_escape(text: str) -> str:
    replacements = {
        "\\": r"\textbackslash{}",
        "&": r"\&",
        "%": r"\%",
        "$": r"\$",
        "#": r"\#",
        "_": r"\_",
        "{": r"\{",
        "}": r"\}",
        "~": r"\textasciitilde{}",
        "^": r"\textasciicircum{}",
    }
    return "".join(replacements.get(ch, ch) for ch in text)


def code_block_tex(code: str) -> str:
    return "\\begin{lstlisting}\n" + code.strip("\n") + "\n\\end{lstlisting}\n"


def paragraph(text: str) -> dict:
    return {"type": "p", "text": text}


def bullets(items: list[str]) -> dict:
    return {"type": "bullets", "items": items}


def code(code_text: str) -> dict:
    return {"type": "code", "code": code_text.strip("\n")}


sections = [
    {
        "title": "Abstract",
        "items": [
            paragraph(
                "This report explains how the Java Pac-Man project applies the main principles of Object-Oriented Programming. "
                "The project is built around reusable game objects, state objects, tile objects, input handlers, file utilities, and rendering helpers. "
                "The report focuses on the concrete OOP concepts used in the code: inheritance, encapsulation, information hiding, abstraction, polymorphism, interface implementation, composition, subtyping, exception handling, customization, and the Open/Closed Principle."
            )
        ],
    },
    {
        "title": "Project Introduction",
        "items": [
            paragraph(
                "The project is a desktop Pac-Man-style arcade game implemented in Java for an Object-Oriented Programming course. "
                "The player controls Pac-Man inside a maze, collects dots, avoids ghosts, moves between levels, controls sound, saves or loads progress, and records scores on a leaderboard. "
                "The goal of the implementation is not only to create a playable game, but also to demonstrate how a medium-size program can be divided into cooperating objects with clear responsibilities."
            ),
            paragraph(
                "The game uses a traditional update-and-draw loop. During each frame, the program reads input, updates the active game state, updates entities, checks collisions, and then draws the current state to the window. "
                "This structure makes OOP useful because many different objects need to react to the same actions, such as update and draw, while keeping their own specialized behavior."
            ),
            code(
                """public class Game implements Runnable {
    private GameWindow wnd;
    private final KeyManager key;
    private final Mouse mouse;
    private State menuState;
    private int level;
    private int score;
    private final GameFileManager gameFileManager;

    @Override
    public void run() {
        InitGame();
        while (runState) {
            Update();
            Draw();
        }
    }
}"""
            ),
        ],
    },
    {
        "title": "Project Overview",
        "items": [
            paragraph(
                "The architecture follows a modular package structure. The central GAME package contains the main game loop, the window class, map loading, level switching, sound, persistence, and shared access through Handler. "
                "The Handler class is important because it gives entities and states controlled access to the current Game, Map, input managers, and GameFileManager without forcing every class to create those objects by itself."
            ),
            paragraph(
                "The entity system represents everything interactive inside the game world. Entity is the abstract base type for all objects that can update, draw, and die. Creature extends Entity for moving objects, while StaticEntity is used for non-moving objects such as collectible dots. "
                "Concrete subclasses such as Player, RedGhost, BlueGhost, PinkGhost, OrangeGhost, and Dot implement the actual game behavior."
            ),
            paragraph(
                "The state system separates the screens of the game. MenuState, GameState, AboutState, SoundState, and LeaderboardState all extend the abstract State class. "
                "This lets the program switch between menu, gameplay, level selection, sound options, and leaderboard without mixing all screen logic in one class."
            ),
            paragraph(
                "The map is tile-based. Map reads a text file, converts numeric tile identifiers into Tile objects, stores them in a generic Grid<Tile>, and draws them on screen. "
                "Concrete tile classes such as Block, DotTile, gate tiles, corner tiles, and directional tiles extend Tile and customize whether the tile is solid and which image it uses."
            ),
            paragraph(
                "Graphics are organized through Assets, ImageLoader, and SpriteSheet. Assets loads and exposes images, ImageLoader handles image loading and resizing, and SpriteSheet extracts smaller images from larger sprite sheets. "
                "Input is handled by KeyManager and Mouse, which implement Java listener interfaces. Persistence is handled by GameFileManager, which saves the current level, lives, score, and leaderboard values."
            ),
            bullets(
                [
                    "GAME: game loop, window, map, levels, sound, file management, and shared Handler.",
                    "GAME.Entities: abstract entity hierarchy and manager for active objects.",
                    "GAME.Entities.MovingCreatures: Player and ghosts, all based on Creature.",
                    "GAME.States: menu, gameplay, level selection, sound settings, and leaderboard screens.",
                    "GAME.Tiles: reusable tile hierarchy for the maze.",
                    "GAME.Graphics and GAME.Keys: image loading, sprite handling, keyboard input, and mouse input.",
                ]
            ),
        ],
    },
    {
        "title": "Inheritance",
        "items": [
            paragraph(
                "Inheritance is used when a class needs to reuse common attributes and behavior from a more general class. "
                "The project uses inheritance heavily in the entity, state, and tile hierarchies. For example, Player is a specialized Creature, and Creature is a specialized Entity."
            ),
            code(
                """public abstract class Entity {
    public abstract void Update() throws InvalidFileException;
    public abstract void Draw(Graphics g);
    public abstract void die();
}

public abstract class Creature extends Entity {
    public Creature(Handler handler, double x, double y) {
        this(handler, x, y, DEFAULT_WIDTH, DEFAULT_HEIGHT);
    }
}

public class Player extends Creature implements Damageable {
    public Player(Handler handler, float x, float y) {
        super(handler, x, y);
    }
}"""
            ),
            paragraph(
                "This hierarchy means Player does not need to reimplement the basic position, size, collision area, and movement structure already defined in Entity and Creature. "
                "The same idea is used by the ghost classes and by tile classes such as Block extending Tile."
            ),
        ],
    },
    {
        "title": "Encapsulation",
        "items": [
            paragraph(
                "Encapsulation combines state and behavior inside a class. The internal fields are kept private, and other classes interact with them through methods. "
                "This protects the object from uncontrolled changes and keeps the rules for changing state inside the class that owns the data."
            ),
            code(
                """public abstract class Entity {
    private double x, y;
    private String direction = "right";
    private final Handler handler;
    private int width, height;
    private boolean active = true;

    public double getX() { return x; }
    public double getY() { return y; }
    public void setx(double x) { this.x = x; }
    public void sety(double y) { this.y = y; }
    public boolean isActive() { return active; }
    public void setActive(boolean active) { this.active = active; }
}"""
            ),
            paragraph(
                "The position and active status of an Entity are not accessed directly from outside. "
                "Classes such as Player, Creature, and EntityManager use getters and setters, which makes the code easier to control and maintain."
            ),
        ],
    },
    {
        "title": "Information Hiding",
        "items": [
            paragraph(
                "Information hiding separates the public interface from the private implementation. The public part shows what other classes can use, while the private part hides data and algorithms that should not be changed directly. "
                "This project demonstrates information hiding through private attributes and private helper methods."
            ),
            code(
                """public class GameWindow {
    private JFrame frame;
    private final String title;
    private final int width;
    private final int height;
    private Canvas canvas;

    public GameWindow(String title, int width, int height) {
        this.title = title;
        this.width = width;
        this.height = height;
        buildGameWindow();
    }

    private void buildGameWindow() {
        frame = new JFrame(title);
        canvas = new Canvas();
        frame.add(canvas);
    }

    public Canvas GetCanvas() { return canvas; }
    public JFrame getFrame() { return frame; }
}"""
            ),
            paragraph(
                "The private attributes frame, title, width, height, and canvas are implementation details of the window. "
                "External classes do not construct or modify the window internals directly. They use public methods such as GetCanvas() and getFrame()."
            ),
            code(
                """public abstract class Creature extends Entity {
    public void move() {
        moveX();
        moveY();
    }

    private void moveX() {
        // horizontal collision and movement algorithm
    }

    private void moveY() {
        // vertical collision and movement algorithm
    }
}"""
            ),
            paragraph(
                "The public method move() is the visible interface. The algorithms moveX() and moveY() are hidden because other classes only need to ask the creature to move, not control every internal collision calculation."
            ),
        ],
    },
    {
        "title": "Abstraction",
        "items": [
            paragraph(
                "Abstraction means defining a general concept and hiding the specific details behind a common interface. "
                "The project abstracts game screens with the State class. Every screen must update and draw, but each screen decides how those operations work."
            ),
            code(
                """public abstract class State {
    private static State currentState = null;
    private final Handler handler;

    public static void setState(State state) {
        State.currentState = state;
    }

    public static State getState() {
        return currentState;
    }

    public abstract void Update() throws InvalidFileException;
    public abstract void Draw(Graphics g);
}"""
            ),
            paragraph(
                "MenuState, GameState, AboutState, SoundState, and LeaderboardState are different concrete states, but the Game class can work with them through the same State abstraction."
            ),
        ],
    },
    {
        "title": "Polymorphism",
        "items": [
            paragraph(
                "Polymorphism allows objects of different concrete classes to be used through a shared type. "
                "This project uses inclusion polymorphism, method overloading, parametric polymorphism, and type coercion or casting."
            ),
        ],
    },
    {
        "title": "Polymorphism by Inclusion",
        "items": [
            paragraph(
                "Polymorphism by inclusion appears when subclasses are treated as objects of a parent type. "
                "EntityManager stores many different objects in one ArrayList<Entity> and calls the same methods on each object."
            ),
            code(
                """public class EntityManager {
    private ArrayList<Entity> entities;

    public void Update() throws InvalidFileException {
        Iterator<Entity> it = entities.iterator();
        while (it.hasNext()) {
            Entity e = it.next();
            e.Update();
            if (!e.isActive()) {
                it.remove();
            }
        }
    }

    public void Draw(Graphics g) {
        for (Entity e : entities) {
            e.Draw(g);
        }
    }
}"""
            ),
            paragraph(
                "At runtime, Java dispatches to the correct Update() or Draw() implementation for Player, Dot, or a ghost. "
                "The manager does not need separate loops for every concrete class."
            ),
        ],
    },
    {
        "title": "Method Overloading",
        "items": [
            paragraph(
                "Method overloading allows multiple methods or constructors to share a name while using different parameter lists. "
                "Creature has two constructors: one with custom dimensions and one with default dimensions."
            ),
            code(
                """public abstract class Creature extends Entity {
    public Creature(Handler handler, double x, double y,
                    int width, int height) {
        super(handler, x, y, width, height);
    }

    public Creature(Handler handler, double x, double y) {
        this(handler, x, y, DEFAULT_WIDTH, DEFAULT_HEIGHT);
    }
}"""
            ),
            paragraph("ImageLoader also overloads LoadImage so an image can be loaded normally or loaded and resized in one call."),
            code(
                """public static BufferedImage LoadImage(String path)
        throws ImageNotFoundException {
    return ImageIO.read(ImageLoader.class.getResource(path));
}

public static BufferedImage LoadImage(String path, int width, int height)
        throws ImageNotFoundException {
    BufferedImage originalImage = LoadImage(path);
    BufferedImage resizedImage =
        new BufferedImage(width, height, originalImage.getType());
    return resizedImage;
}"""
            ),
        ],
    },
    {
        "title": "Parametric Polymorphism",
        "items": [
            paragraph(
                "Parametric polymorphism is implemented with Java generics. Grid<T> can store any type T while still keeping type safety."
            ),
            code(
                """public class Grid<T> {
    private final List<List<T>> grid;

    public void setCell(int row, int col, T item) {
        if (isValid(row, col)) {
            grid.get(row).set(col, item);
        }
    }

    public T getCell(int row, int col) {
        if (isValid(row, col)) {
            return grid.get(row).get(col);
        }
        return null;
    }
}"""
            ),
            code(
                """private Grid<Tile> gameTiles;
gameTiles = new Grid<>(height, width);"""
            ),
            paragraph(
                "The Map class uses Grid<Tile>, so the grid is reusable but still restricted to Tile objects in this context."
            ),
        ],
    },
    {
        "title": "Polymorphism by Coercion and Type Casting",
        "items": [
            paragraph(
                "Coercion and casting allow a value to be treated as another compatible type. "
                "The project uses numeric casts for drawing and runtime type checks for collision behavior."
            ),
            code(
                """public void Draw(Graphics g) {
    BufferedImage image = directionImages.get(getDirection());
    g.drawImage(image, (int) getX(), (int) getY(), null);
}"""
            ),
            code(
                """for (Entity e : getHandler().getMap()
        .getEntityManager().getEntities()) {
    if (e instanceof Dot) {
        e.setActive(false);
        getHandler().getGame().setScore(10);
    }
}"""
            ),
            paragraph(
                "The cast from double to int is needed because drawing coordinates must be integers. "
                "The instanceof check lets the program apply dot-specific logic only when the collided Entity is actually a Dot."
            ),
        ],
    },
    {
        "title": "Interface Implementation",
        "items": [
            paragraph(
                "An interface defines a contract that a class must follow. The project defines Damageable for objects that can receive damage."
            ),
            code(
                """public interface Damageable {
    void takeDamage(int amount);
}

public class Player extends Creature implements Damageable {
    @Override
    public void takeDamage(int amount) {
        System.out.println("Player taking " + amount + " damage!");
        die();
    }
}"""
            ),
            paragraph("The project also uses Java event interfaces for input handling."),
            code(
                """public class KeyManager implements KeyListener {
    @Override
    public void keyPressed(KeyEvent e) {
        keys[e.getKeyCode()] = true;
    }
}

public class Mouse implements MouseListener, MouseMotionListener {
    @Override
    public void mouseMoved(MouseEvent e) {
        x = e.getX();
        y = e.getY();
    }
}"""
            ),
        ],
    },
    {
        "title": "Composition",
        "items": [
            paragraph(
                "Composition is a has-a relationship. A complex object is built from smaller objects that each have their own responsibility. "
                "The Game class is composed of a window, keyboard manager, mouse manager, state references, and a file manager."
            ),
            code(
                """public class Game implements Runnable {
    private GameWindow wnd;
    private final KeyManager key;
    private final Mouse mouse;
    private State menuState;
    private final GameFileManager gameFileManager;

    public Game(String title, int width, int intHeight,
                GameFileManager gameFileManager) {
        key = new KeyManager();
        mouse = new Mouse();
        this.gameFileManager = gameFileManager;
    }
}"""
            ),
            paragraph("The Map class is also composed of a grid and an entity manager."),
            code(
                """public class Map {
    private Grid<Tile> gameTiles;
    private final EntityManager entityManager;

    public Map(Handler handler, String path)
            throws InvalidFileException {
        entityManager =
            new EntityManager(handler, new Player(handler, 100, 100));
        loadMap(path);
    }
}"""
            ),
        ],
    },
    {
        "title": "Subtyping",
        "items": [
            paragraph(
                "Subtyping means a derived type can be used where a base type or interface is expected. "
                "The developer side of subtyping is visible in the class declarations, where the relationships are created."
            ),
            code(
                """public class Player extends Creature implements Damageable {
    // Player is a Creature, an Entity, and a Damageable object.
}

public class RedGhost extends Creature {
    // RedGhost is a Creature and an Entity.
}

public class MenuState extends State {
    // MenuState is a State.
}

public class Block extends Tile {
    // Block is a Tile.
}"""
            ),
            paragraph(
                "These declarations are the developer-side proof of subtyping. They create the relationships that later allow generic entity lists and state switching."
            ),
        ],
    },
    {
        "title": "Exception Handling",
        "items": [
            paragraph(
                "Exception handling lets the program react to invalid files, missing resources, malformed data, and other runtime problems without mixing error logic into normal game logic. "
                "The project uses Java built-in exceptions and custom exceptions."
            ),
            code(
                """public class InvalidFileException extends Exception {
    InvalidFileException(String message) {
        super(message);
    }
}

public class ImageNotFoundException extends Exception {
    ImageNotFoundException(String message) {
        super(message);
    }
}"""
            ),
            code(
                """public static String loadFile(String path)
        throws InvalidFileException {
    StringBuilder builder = new StringBuilder();
    try {
        BufferedReader b = new BufferedReader(new FileReader(path));
        String line;
        while ((line = b.readLine()) != null) {
            builder.append(line).append("\\n");
        }
        b.close();
    } catch (Exception e) {
        throw new InvalidFileException("The file was not found!");
    }
    return builder.toString();
}"""
            ),
            paragraph(
                "GameFileManager catches errors when reading saved progress. If the save file is missing or invalid, it falls back to a new game state."
            ),
        ],
    },
    {
        "title": "Customization",
        "items": [
            paragraph(
                "Customization in this project means behavior that is not fixed to one hard-coded path and can change based on user choice or saved user data. "
                "The current project does not implement choosing a different player character, but it does support level selection, sound choice, save/load progress, and leaderboard persistence."
            ),
            paragraph("The player can choose the level from AboutState. The selected level changes the active GameState and saved game data."),
            code(
                """if (getHandler().getMouse().pressed_left()
        && level4.contains(getHandler().getMouse().getX(),
                           getHandler().getMouse().getY())) {
    getHandler().getGame().setLevel(4);
    getHandler().getGameFileManager().saveGame(4, 5, 0);
    State.setState(new GameState(getHandler(), 4));
}"""
            ),
            paragraph("The player can load previous progress from the main menu."),
            code(
                """if (getHandler().getMouse().pressed_left()
        && load.contains(getHandler().getMouse().getX(),
                         getHandler().getMouse().getY())) {
    int[] savedData = getHandler().getGameFileManager().loadGame();
    getHandler().getGame().setLevel(savedData[0]);
    getHandler().getGame().setScoreValue(savedData[2]);
    State.setState(new GameState(getHandler(), savedData[0]));
}"""
            ),
            paragraph("The sound screen lets the user turn off the game sounds."),
            code(
                """if (getHandler().getMouse().pressed_left()
        && yes.contains(getHandler().getMouse().getX(),
                        getHandler().getMouse().getY())) {
    Sound.getBeginningSound().setSoundOff(1);
    Sound.getChompSound().setSoundOff(1);
    Sound.getDeathSound().setSoundOff(1);
    State.setState(getHandler().getGame().getMenuState());
}"""
            ),
            paragraph(
                "Save/load and leaderboard files also personalize the game because the result depends on the user's previous progress and scores instead of a single fixed initial state."
            ),
        ],
    },
    {
        "title": "Open/Closed Principle",
        "items": [
            paragraph(
                "The Open/Closed Principle states that classes should be open for extension but closed for modification. "
                "The project supports this idea through abstract base classes and polymorphic loops. New states, entities, and tiles can be added by extending the existing base types."
            ),
            code(
                """public class Block extends Tile {
    public Block(int idd) {
        super(Assets.getBlock(), idd);
    }

    @Override
    public boolean IsSolid() {
        return true;
    }
}"""
            ),
            code(
                """for (Entity e : entities) {
    e.Draw(g);
}

if (State.getState() != null) {
    State.getState().Update();
}"""
            ),
            paragraph(
                "The loops work with Entity and State references, so adding a new enemy or screen does not require rewriting the main draw/update mechanism. "
                "The tile hierarchy works similarly: each new tile can override behavior such as IsSolid() while the map still works with the general Tile type."
            ),
        ],
    },
    {
        "title": "Conclusion",
        "items": [
            paragraph(
                "The Pac-Man project demonstrates Object-Oriented Programming through a practical game architecture. "
                "Inheritance, abstraction, and subtyping organize common concepts; encapsulation and information hiding protect internal details; polymorphism lets the game loop work with many concrete objects; interfaces define behavioral contracts; and composition builds the application from focused components. "
                "The customization features and Open/Closed structure also show how the project can support user choices and future extension."
            )
        ],
    },
]


def build_tex(path: Path) -> None:
    parts: list[str] = [
        r"\documentclass[12pt,a4paper]{article}",
        r"\usepackage[margin=2.5cm]{geometry}",
        r"\usepackage[T1]{fontenc}",
        r"\usepackage[utf8]{inputenc}",
        r"\usepackage{lmodern}",
        r"\usepackage{xcolor}",
        r"\usepackage{hyperref}",
        r"\usepackage{listings}",
        r"\usepackage{setspace}",
        r"\usepackage{titlesec}",
        r"\usepackage{titling}",
        r"\setstretch{1.15}",
        r"\definecolor{codebg}{RGB}{247,248,250}",
        r"\definecolor{codeborder}{RGB}{205,213,223}",
        r"\definecolor{codekw}{RGB}{27,83,156}",
        r"\definecolor{codestr}{RGB}{139,79,18}",
        r"\definecolor{codecomment}{RGB}{82,120,82}",
        r"\lstdefinestyle{javastyle}{language=Java,basicstyle=\ttfamily\small,keywordstyle=\color{codekw}\bfseries,stringstyle=\color{codestr},commentstyle=\color{codecomment},backgroundcolor=\color{codebg},frame=single,rulecolor=\color{codeborder},framerule=0.5pt,breaklines=true,columns=fullflexible,showstringspaces=false,tabsize=4,xleftmargin=0.25cm,xrightmargin=0.25cm,aboveskip=10pt,belowskip=10pt}",
        r"\lstset{style=javastyle}",
        r"\titleformat{\section}{\Large\bfseries\color{black}}{}{0pt}{}",
        r"\titleformat{\subsection}{\large\bfseries\color{black}}{}{0pt}{}",
        r"\hypersetup{colorlinks=true,linkcolor=black,urlcolor=black}",
        r"\begin{document}",
        r"\begin{titlepage}",
        r"\centering",
        r"{\large Department MIFT\par}",
        r"\vspace{0.4cm}",
        r"{\large Course: Object-Oriented Programming\par}",
        r"\vspace{1.5cm}",
        r"{\Huge\bfseries PAC-MAN JAVA GAME PROJECT\par}",
        r"\vspace{0.8cm}",
        r"{\Large Object-Oriented Programming Report\par}",
        r"\vfill",
        r"{\large Student: Mehdi Talebikhatir\par}",
        r"{\large Matricola: 558948\par}",
        r"\vspace{0.4cm}",
        r"{\large Professor: Salvatore Distefano\par}",
        r"{\large Universita degli Studi di Messina\par}",
        r"\vfill",
        r"\end{titlepage}",
        r"\tableofcontents",
        r"\newpage",
    ]
    for sec in sections:
        title = latex_escape(sec["title"])
        parts.append(rf"\section*{{{title}}}")
        parts.append(rf"\addcontentsline{{toc}}{{section}}{{{title}}}")
        for item in sec["items"]:
            if item["type"] == "p":
                parts.append(latex_escape(item["text"]) + "\n")
            elif item["type"] == "bullets":
                parts.append(r"\begin{itemize}")
                for bullet in item["items"]:
                    parts.append(r"\item " + latex_escape(bullet))
                parts.append(r"\end{itemize}")
            elif item["type"] == "code":
                parts.append(code_block_tex(item["code"]))
        parts.append("")
    parts.append(r"\end{document}")
    path.write_text("\n".join(parts), encoding="utf-8")


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_border(cell, color: str = "D0D7DE") -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = "w:{}".format(edge)
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def add_code_block_docx(doc: Document, code_text: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.autofit = True
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F6F8FA")
    set_cell_border(cell)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(code_text.strip("\n"))
    run.font.name = "Consolas"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(35, 39, 47)
    doc.add_paragraph()


def add_footer(section) -> None:
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("Pac-Man Java Game Project - OOP Report")
    run.font.name = "Calibri"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(100, 100, 100)


def build_docx(path: Path) -> None:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    add_footer(section)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for style_name, size, color in [
        ("Heading 1", 16, RGBColor(31, 78, 121)),
        ("Heading 2", 13, RGBColor(31, 78, 121)),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(6)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("PAC-MAN JAVA GAME PROJECT")
    r.bold = True
    r.font.size = Pt(24)
    r.font.name = "Calibri"
    r.font.color.rgb = RGBColor(31, 78, 121)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = subtitle.add_run("Object-Oriented Programming Report")
    r.font.size = Pt(15)
    r.font.name = "Calibri"

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(
        "Department MIFT\n"
        "Course: Object-Oriented Programming\n\n"
        "Student: Mehdi Talebikhatir\n"
        "Matricola: 558948\n"
        "Professor: Salvatore Distefano\n"
        "Universita degli Studi di Messina"
    )
    doc.add_page_break()

    for sec in sections:
        doc.add_heading(sec["title"], level=1)
        for item in sec["items"]:
            if item["type"] == "p":
                doc.add_paragraph(item["text"])
            elif item["type"] == "bullets":
                for bullet in item["items"]:
                    doc.add_paragraph(bullet, style="List Bullet")
            elif item["type"] == "code":
                add_code_block_docx(doc, item["code"])

    doc.save(path)


def unique_paths() -> tuple[Path, Path, Path]:
    REPORTS.mkdir(exist_ok=True)
    stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    base = f"OOP_Project_Report_{stamp}"
    tex = REPORTS / f"{base}.tex"
    docx = REPORTS / f"{base}.docx"
    pdf = REPORTS / f"{base}.pdf"
    n = 2
    while tex.exists() or docx.exists() or pdf.exists():
        base_n = f"{base}_{n}"
        tex = REPORTS / f"{base_n}.tex"
        docx = REPORTS / f"{base_n}.docx"
        pdf = REPORTS / f"{base_n}.pdf"
        n += 1
    return tex, docx, pdf


if __name__ == "__main__":
    tex_path, docx_path, pdf_path = unique_paths()
    build_tex(tex_path)
    build_docx(docx_path)
    print(tex_path)
    print(docx_path)
    print(pdf_path)
